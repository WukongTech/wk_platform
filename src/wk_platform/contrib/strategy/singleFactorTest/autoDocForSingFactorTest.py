# -*- coding: utf-8 -*-
"""
Created on Tue Jan 23 14:41:02 2018

@author: yangr
"""

import sys 
"""
reload(sys) 
sys.setdefaultencoding('utf8')
"""
from docx import Document
from docx.shared import Inches

import matplotlib as mpl

import matplotlib.pyplot as plt

import seaborn as sns
import pandas as pd
import numpy as np
import datetime
import os
import shutil
from matplotlib.ticker import FuncFormatter, MaxNLocator
from docx.oxml.ns import qn
import copy
from docx.shared import Pt
from docx import styles
from docx.enum.text import WD_ALIGN_PARAGRAPH



class portfolioAutoTrackDoc():
    def __init__(self):
        """
        创建默认文档
        """
        self.__document = Document()
        self.__document.styles['Normal'].font.name = '楷体'
        self.__document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False
        
        
    def documentFrame(self,factorName):
        self.__document.add_heading('悟空单因子测试报告%s'%(factorName), level = 0)
    
    
    def __backTestCompare(self,factorName):
        
        sheetName  = '指数净值'
        netValue_ALL = pd.DataFrame()
        
        for groupNum in range(1,6):
            backTestResult = pd.read_excel('dataFile/result/20050601_20170914_backtest_weightFinalNotEqualOrg__equal_MA_1_%s_group%s.csv.xlsx'%(factorName,str(groupNum)), sheetName)
            netValue = backTestResult['策略对冲沪深300']
            netValue_ALL['Date'] = backTestResult['日期']
            netValue_ALL['group'+str(groupNum)] = netValue
        netValue_ALL['minus'] = netValue_ALL['group5']-netValue_ALL['group1']
        #print netValue_ALL
        x=[datetime.datetime.strptime(d, '%Y-%m-%d').date() for d in netValue_ALL['Date']]
        plt.plot(x,netValue_ALL['minus'])
        plt.savefig('Plots/' + factorName + '_BackTestPlots_curve.png') 
        self.__document.add_heading('分层回测多空净值曲线', level = 3)
        pngName ='Plots/'+ factorName + '_BackTestPlots_curve.png' 
        
        self.__document.add_picture(pngName,width=Inches(4))
        plt.show() 
        
        sheetName  = '策略指标'
        strategyReturn_ALL = []
        strategySharpe_ALL = []
        strategy_index_ALL = pd.DataFrame(columns = ['策略收益','年化收益率','波动率','最大回撤','夏普比率','月胜率','盈亏比'])
        for groupNum in range(1,6):
            backTestResult = pd.read_excel('dataFile/result/20050601_20170914_backtest_weightFinalNotEqualOrg__equal_MA_1_%s_group%s.csv.xlsx'%(factorName,str(groupNum)), sheetName)
           
            strategyReturn = backTestResult.loc['策略对冲沪深300','年化收益率']
            
            strategy_index = backTestResult.loc['策略对冲沪深300',['策略收益','年化收益率','波动率','最大回撤','夏普比率','月胜率','盈亏比']]
            strategy_indexList = strategy_index.tolist()
            print(strategy_indexList[0])
            strategy_index = pd.DataFrame([[strategy_indexList[0],strategy_indexList[1],strategy_indexList[2],strategy_indexList[3],strategy_indexList[4],strategy_indexList[5],strategy_indexList[6]]],columns = ['策略收益','年化收益率','波动率','最大回撤','夏普比率','月胜率','盈亏比'])
            strategy_index.rename({0:'group'+str(groupNum)},inplace = True)
            print(strategy_index)
            strategySharpe = backTestResult.loc['策略对冲沪深300','夏普比率']
            strategySharpe_ALL.append(strategySharpe)
            
            strategyReturn_ALL.append(strategyReturn)
            strategy_index_ALL = strategy_index_ALL.append(strategy_index)
          
            
        print(strategyReturn_ALL)
        print(strategy_index_ALL)
        plt.bar(list(range(0,5)),strategyReturn_ALL)
        plt.xticks(list(range(0,5)),['group'+str(i) for i in range(1,6)])
     
        plt.savefig('Plots/' + factorName + '_BackTestPlots_return.png')
        self.__document.add_heading('分层回测分组年化收益率', level = 3)
        pngName ='Plots/'+ factorName + '_BackTestPlots_return.png' 
        
        self.__document.add_picture(pngName,width=Inches(4))
        
        plt.bar(list(range(0,5)),strategySharpe_ALL)
        plt.xticks(list(range(0,5)),['group'+str(i) for i in range(1,6)])
        plt.savefig('Plots/' + factorName + '_BackTestPlots_sharpe.png')
        
        self.__document.add_heading('分层回测分组夏普比率', level = 3)
        pngName ='Plots/'+ factorName + '_BackTestPlots_sharpe.png' 
        
        self.__document.add_picture(pngName,width=Inches(4))
        
     
        self.__document.add_heading('分层回测分组策略指标', level = 3)
        data = strategy_index_ALL
        columns = ['年化收益率','波动率',	'最大回撤',	'夏普比率',	'月胜率',	'盈亏比']
        data = data[columns]
        data = data.round(4)
        """
        添加表
        """
        rowNumber = data.shape[0]
        colNumber = data.shape[1]
        table = self.__document.add_table(rowNumber+1, colNumber+1)
        """
        print data.columns
        print 'get number'
        print rowNumber
        print len(data.index)
        """
        rowCount = 0
        ColCount = 1
        for column in data.columns:
            row = table.rows[rowCount].cells
            row[ColCount].text = column
            ColCount +=1          
            
        rowCount = 1
        ColCount = 0
        for indexs in data.index:
            print(indexs)
            ColCount = 0
            row = table.rows[rowCount].cells
           
            if ColCount == 0:
                row[ColCount].text = indexs
                ColCount +=1
            for column in data.columns:
    
                res = (data.loc[indexs][column])
                res = '%.2f%%' % (res * 100)
                row[ColCount].text = str(res)
                #row[ColCount].text = str(data.loc[indexs][column])
                ColCount +=1
            rowCount += 1
                
        table.style = 'LightShading-Accent1'
        
        
    """
    整个文档的布局
    """
    def dailyStatictics(self,factorName, benchIndex):
        
        sheetName = '单因子测试'
        

        
        self.__document.add_paragraph('测试区间：20050408-20170703',style = 'ListBullet')
        
        self.__document.add_heading('一.因子的含义和来源', level = 1)
        """
        输入因子含义
        """
        self.__document.add_heading('二.单因子回归测试和IC值结果展示', level = 1)
        self.__document.add_heading('2.1因子回归测试和IC值', level = 2)
       
        
        singleFactor = pd.read_excel('dataFile/PART1/singleFactor/singleFactorOutput_%s_.xlsx'%(factorName), sheetName)
        #style='ListBullet'
        rowNumber = singleFactor.shape[0]
        colNumber = singleFactor.shape[1]-1
        print(rowNumber)
        print(colNumber)
        table = self.__document.add_table(rowNumber+1, colNumber)
        styles.style._TableStyle.font = Pt(9)
        rowCount = 0
        colCount = 0
        """
        记录列名
        """
        for column in singleFactor.columns[1:]:
            row = table.rows[rowCount].cells
            row[colCount].text = column
            colCount += 1
        rowCount +=1
        colCount = 0
    
        """
        记录表格内容
        """
        print(singleFactor.index)
        for indexs in singleFactor.index:
            for column in singleFactor.columns[1:]:
                row = table.rows[rowCount].cells             
                if column != '因子名称':    
                 
                    #print (daysChange.loc[indexs][column])
                    #print type((daysChange.loc[indexs][column]))
                    res = float(singleFactor.loc[indexs][column]) 
                    res = '%.1f%%' % (res * 100)
                    row[colCount].text = str(res)
                else:
                    row[colCount].text = str(singleFactor.loc[indexs][column])
                colCount += 1
            colCount = 0
            rowCount +=1
        table.style = 'LightGrid-Accent1'
        
          
    
        
        
        
        self.__document.add_heading('2.2因子收益率和IC值分布图', level = 2)
        
 
        pngName = 'Plots/'+factorName+'_ _AnalysisPlots.png'
  
        self.__document.add_picture(pngName,width = Inches(4.2))
        
        
        
        self.__document.add_heading('2.3因子分层回测结果', level = 2)
        
        self.__document.add_heading('2.3.1 全市场分层回测', level = 3)
        self.__backTestCompare(factorName)
        
        '''
        self.__document.add_heading(u'2.3.2 指数单因子策略回测', level = 3)
        sheetNameList = [u'策略绝对收益', u'策略相对收益']
        
        self.__document.add_heading(sheetName, level = 3)
        data= pd.read_excel(u'dataFile/result/20050601_20170914_backtest_weightFinalNotEqualOrg__equal_MA_1_%s.csv.xlsx'%(factorName),u'策略指标',encoding = 'gbk')
        columns = [u'年化收益率',u'波动率',	u'最大回撤',	u'夏普比率',	u'月胜率',	u'盈亏比']
        data = data[columns]
        data = data.round(4)
        """
        添加表
        """
        rowNumber = data.shape[0]
        colNumber = data.shape[1]
        table = self.__document.add_table(rowNumber+1, colNumber+1)
        """
        print data.columns
        print 'get number'
        print rowNumber
        print len(data.index)
        """
        rowCount = 0
        ColCount = 1
        for column in data.columns:
            row = table.rows[rowCount].cells
            row[ColCount].text = column
            ColCount +=1          
            
        rowCount = 1
        ColCount = 0
        for indexs in data.index:
            print indexs
            ColCount = 0
            row = table.rows[rowCount].cells
            if ColCount == 0:
                row[ColCount].text = indexs
                ColCount +=1
            for column in data.columns:
    
                res = (data.loc[indexs][column])
                res = '%.2f%%' % (res * 100)
                row[ColCount].text = str(res)
                #row[ColCount].text = str(data.loc[indexs][column])
                ColCount +=1
            rowCount += 1
                
        table.style = 'LightShading-Accent1'
        '''
        self.__document.add_page_break()
      
        self.__document.add_heading('2.3不同市场的因子表现', level = 2)




        sheetNameList = ['牛市', '熊市','震荡市','最近一年']
        for sheetName in sheetNameList:
            if sheetName == '牛市':
                self.__document.add_heading("牛市", level = 3)
                #self.__document.add_paragraph(u"牛市",style = 'ListBullet')
                
                self.__document.add_paragraph("2005-07-11-2007-10-16")
                self.__document.add_paragraph("2008-11-04-2009-08-04")
                self.__document.add_paragraph("2014-05-31-2015-06-12")
                self.__document.add_paragraph("2015-08-26-2015-12-22")
                self.__document.add_paragraph("2016-01-28-2017-08-28")
            elif sheetName == '熊市':
                self.__document.add_heading("熊市", level = 3)
               
                #self.__document.add_paragraph(u"熊市",style = 'ListBullet')
                self.__document.add_paragraph("2007-10-16-2008-11-04")
                self.__document.add_paragraph("2010-11-08-2012-12-13")
                self.__document.add_paragraph("2015-06-12-2015-08-26")
            elif sheetName == '震荡市':
                self.__document.add_heading("震荡市", level = 3)
           
                #self.__document.add_paragraph(u"震荡市",style = 'ListBullet')
                self.__document.add_paragraph("2009-08-04-2010-11-08")
                self.__document.add_paragraph("2012-12-03-2014-05-31")
            else:
                self.__document.add_heading(sheetName, level = 3)
            
            singleFactor = pd.read_excel('dataFile/PART1/singleFactor/singleFactorOutput_%s_.xlsx'%(factorName), sheetName ,encoding = 'gbk')
            
            rowNumber = singleFactor.shape[0]
            colNumber = singleFactor.shape[1]
            if sheetName == '最近一年':
                table = self.__document.add_table(rowNumber+1, colNumber)
            else:
                table = self.__document.add_table(rowNumber+1, colNumber-1)
            rowCount = 0
            colCount = 0
            """
            记录列名
            """
            if sheetName == '最近一年':
                columnList = singleFactor.columns
            else:
                columnList = singleFactor.columns[:-1]
            for column in columnList:
                row = table.rows[rowCount].cells
                row[colCount].text = column
                colCount += 1
            rowCount +=1
            colCount = 0
        
            """
            记录表格内容
            """
            
            for indexs in singleFactor.index:
                for column in columnList:
                    row = table.rows[rowCount].cells             
                    if (column != '因子名称') and (column != 'Time'):    
                     
                        #print (daysChange.loc[indexs][column])
                        #print type((daysChange.loc[indexs][column]))
                        res = float(singleFactor.loc[indexs][column]) 
                        res = '%.1f%%' % (res * 100)
                        row[colCount].text = str(res)
                    else:
                        row[colCount].text = str(singleFactor.loc[indexs][column])
                    colCount += 1
                colCount = 0
                rowCount +=1
            table.style = 'LightGrid-Accent1'
            
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(8.5)
            self.__document.add_paragraph('    ')
            
            
        self.__document.add_heading('2.4最近一年因子收益率和IC值分布图', level = 2)
        pngName ='Plots/'+factorName+'_recent_1yr _AnalysisPlots.png' 
        
        self.__document.add_picture(pngName,width=Inches(4.2))
        
        
        
        
        self.__document.add_heading('三.因子分组特征', level = 1)
        
        self.__document.add_heading('3.1市值分组', level = 2)
        singleFactor = pd.read_csv('dataFile/PART1/singleFactor/singleFactorOutput_MVclass.csv',encoding  = 'gbk')
      
        rowNumber = 3
        colNumber = singleFactor.shape[1]
     
        table = self.__document.add_table(rowNumber+1, colNumber-1)
       
        rowCount = 0
        colCount = 0
        """
        记录列名
        """
        for column in singleFactor.columns[1:]:
            row = table.rows[rowCount].cells
            row[colCount].text = column
            colCount += 1
        rowCount +=1
        colCount = 0
    
        """
        记录表格内容
        """
       
        
        for indexs in singleFactor.index:
            for column in singleFactor.columns[1:]:
                print(singleFactor.loc[indexs][0])
                if singleFactor.loc[indexs][0] == factorName:
                    
                    row = table.rows[rowCount].cells             
                    if column != 'MV' and column!= '平均样本量':    
                     
                        #print (daysChange.loc[indexs][column])
                        #print type((daysChange.loc[indexs][column]))
                        res = float(singleFactor.loc[indexs][column]) 
                        res = '%.1f%%' % (res * 100)
                        row[colCount].text = str(res)
                    elif column == 'MV':
                        row[colCount].text = str(singleFactor.loc[indexs][column])
                    else:
                        res = int(singleFactor.loc[indexs][column]) 
                        
                        row[colCount].text = str(res)
                        
                    colCount += 1
            colCount = 0
            rowCount +=1
        table.style = 'LightGrid-Accent1'
        for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(8)
         
        
        
        self.__document.add_paragraph('    ')
        self.__document.add_paragraph('    ')
        
        self.__document.add_heading('3.2 ROE分组', level = 2)
        singleFactor = pd.read_csv('dataFile/PART1/singleFactor/singleFactorOutput_ROEclass.csv',encoding  = 'gbk')
      
        rowNumber = 3
        colNumber = singleFactor.shape[1]
      
        table = self.__document.add_table(rowNumber+1, colNumber-1)
       
        rowCount = 0
        colCount = 0
        """
        记录列名
        """
        for column in singleFactor.columns[1:]:
            row = table.rows[rowCount].cells
            row[colCount].text = column
            colCount += 1
        rowCount +=1
        colCount = 0
    
        """
        记录表格内容
        """
        print(singleFactor.index)
        for indexs in singleFactor.index:
            for column in singleFactor.columns[1:]:
                if singleFactor.loc[indexs][0] == factorName:
                    row = table.rows[rowCount].cells             
                    if column != 'ROE_TTM' and column!= '平均样本量':    
                     
                        #print (daysChange.loc[indexs][column])
                        #print type((daysChange.loc[indexs][column]))
                        res = float(singleFactor.loc[indexs][column]) 
                        res = '%.1f%%' % (res * 100)
                        row[colCount].text = str(res)
                   
                    elif column == 'ROE_TTM':
                        row[colCount].text = str(singleFactor.loc[indexs][column])
                    else:
                        res = int(singleFactor.loc[indexs][column]) 
                        
                        row[colCount].text = str(res)
                    colCount += 1
            colCount = 0
            rowCount +=1
        table.style = 'LightGrid-Accent1'
        
        for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(8)
                            
        self.__document.add_paragraph('    ')
        
        
        
        self.__document.add_heading('3.3 周期/非周期性行业分组', level = 2)
        singleFactor = pd.read_csv('dataFile/PART1/singleFactor/singleFactorOutput_cycle.csv',encoding  = 'gbk')
      
        rowNumber = 2
        colNumber = singleFactor.shape[1]
     
        table = self.__document.add_table(rowNumber+1, colNumber-1)
       
        rowCount = 0
        colCount = 0
        """
        记录列名
        """
        for column in singleFactor.columns[1:]:
            row = table.rows[rowCount].cells
            row[colCount].text = column
            colCount += 1
        rowCount +=1
        colCount = 0
    
        """
        记录表格内容
        """
        print(singleFactor.index)
        
        for indexs in singleFactor.index:
            for column in singleFactor.columns[1:]:
                print(singleFactor.loc[indexs][0])
                if singleFactor.loc[indexs][0] == factorName:
                    
                    row = table.rows[rowCount].cells             
                    if column != 'cycle' and column!= '平均样本量':    
                     
                        #print (daysChange.loc[indexs][column])
                        #print type((daysChange.loc[indexs][column]))
                        res = float(singleFactor.loc[indexs][column]) 
                        res = '%.1f%%' % (res * 100)
                        row[colCount].text = str(res)
                    elif column == 'cycle':
                        row[colCount].text = singleFactor.loc[indexs][column]
                    else:
                        row[colCount].text = str(int(singleFactor.loc[indexs][column]))
                    colCount += 1
            colCount = 0
            rowCount +=1
        table.style = 'LightGrid-Accent1'
        for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(8)
         
        
        
        self.__document.add_paragraph('    ')
        
        '''
        self.__document.add_heading(u'四.因子在指数中的表现', level = 1)
        

        self.__document.add_heading(u'4.1因子收益率和IC值', level = 2)
        
        
        sheetNameList = [u'单因子测试',u'牛市', u'熊市',u'震荡市',u'最近一年']
        for sheetName in sheetNameList:
            
            if sheetName == u'牛市':
                self.__document.add_heading(u"牛市", level = 3)
                #self.__document.add_paragraph(u"牛市",style = 'ListBullet')
                
                self.__document.add_paragraph(u"2005-07-11-2007-10-16")
                self.__document.add_paragraph(u"2008-11-04-2009-08-04")
                self.__document.add_paragraph(u"2014-05-31-2015-06-12")
                self.__document.add_paragraph(u"2015-08-26-2015-12-22")
                self.__document.add_paragraph(u"2016-01-28-2017-08-28")
            elif sheetName == u'熊市':
                self.__document.add_heading(u"熊市", level = 3)
               
                #self.__document.add_paragraph(u"熊市",style = 'ListBullet')
                self.__document.add_paragraph("2007-10-16-2008-11-04")
                self.__document.add_paragraph("2010-11-08-2012-12-13")
                self.__document.add_paragraph("2015-06-12-2015-08-26")
            elif sheetName == u'震荡市':
                self.__document.add_heading(u"震荡市", level = 3)
           
                #self.__document.add_paragraph(u"震荡市",style = 'ListBullet')
                self.__document.add_paragraph("2009-08-04-2010-11-08")
                self.__document.add_paragraph("2012-12-03-2014-05-31")
                
            else:
                self.__document.add_heading(sheetName, level = 3)
              
            singleFactor = pd.read_excel(u'dataFile/PART1/singleFactor/singleFactorOutput_%s_%s.xlsx'%(factorName,benchIndex), sheetName ,encoding = 'gbk')
            rowNumber = singleFactor.shape[0]
            colNumber = singleFactor.shape[1]
            if (sheetName == u'单因子测试') or (sheetName == u'最近一年'):
                table = self.__document.add_table(rowNumber+1, colNumber)
            else:
                table = self.__document.add_table(rowNumber+1, colNumber-1)
            rowCount = 0
            colCount = 0
            """
            记录列名
            """
            if (sheetName == u'单因子测试') or (sheetName == u'最近一年'):
                columnList = singleFactor.columns
            else:
                columnList = singleFactor.columns[:-1]
                
            for column in columnList:
                row = table.rows[rowCount].cells
                row[colCount].text = column
                colCount += 1
            rowCount +=1
            colCount = 0
        
            """
            记录表格内容
            """
            
            for indexs in singleFactor.index:
                for column in columnList:
                    row = table.rows[rowCount].cells             
                    if (column != u'因子名称') and (column != 'Time'):    
                     
                        #print (daysChange.loc[indexs][column])
                        #print type((daysChange.loc[indexs][column]))
                        res = float(singleFactor.loc[indexs][column]) 
                        res = '%.1f%%' % (res * 100)
                        row[colCount].text = str(res)
                    else:
                        row[colCount].text = str(singleFactor.loc[indexs][column])
                    colCount += 1
                colCount = 0
                rowCount +=1
            table.style = 'LightGrid-Accent1'
            
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(8.5)
            self.__document.add_paragraph(u'    ')
    
        '''
  
    """
    文档保存
    """
    def saveDoc(self,factorName):
        self.__document.save('悟空单因子分析报告_%s.docx'%(factorName))






def main():
    
    autoDoc = portfolioAutoTrackDoc()
    
    factorName = 'OPPToMV'
    
    benchIndex = ''
    autoDoc.documentFrame(factorName)
    
    autoDoc.dailyStatictics(factorName, benchIndex)  
    

    #autoDoc.backTestCompare(factorName)
    autoDoc.saveDoc(factorName)
            

if __name__ == "__main__":
    main()

